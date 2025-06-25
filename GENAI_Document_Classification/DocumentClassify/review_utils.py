import os
import pytesseract
from PIL import Image
from PyPDF2 import PdfReader
from docx import Document
import json
import uuid
from datetime import datetime
from embedding_utils import classify_document

# Directories
DATA_DIR = "data"
UPLOAD_DIR = "uploads"
LOG_FILE = "document_log.json"

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

def _get_data_file():
    return os.path.join(DATA_DIR, "documents.json")

def _load_data():
    if not os.path.exists(_get_data_file()):
        return []
    with open(_get_data_file(), "r") as f:
        return json.load(f)

def _save_data(data):
    with open(_get_data_file(), "w") as f:
        json.dump(data, f, indent=2)

# ✅ Process uploaded document
def process_document(file):
    data = _load_data()

    if any(doc["name"] == file.name for doc in data):
        return

    doc_id = str(uuid.uuid4())
    file_ext = file.name.split('.')[-1].lower()
    save_path = os.path.join(UPLOAD_DIR, f"{doc_id}.{file_ext}")

    with open(save_path, "wb") as f:
        f.write(file.read())

    doc_record = {
        "id": doc_id,
        "name": file.name,
        "path": save_path,
        "status": "Ingested",
        "last_updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

    data.append(doc_record)
    _save_data(data)

    run_ocr_and_classify(save_path)

# ✅ OCR + Gemini classification
def run_ocr_and_classify(file_path):
    ext = file_path.split('.')[-1].lower()
    text = ""
    from time import time
    start = time()

    try:
        if ext in ['png', 'jpg', 'jpeg']:
            text = pytesseract.image_to_string(Image.open(file_path))
        elif ext == 'pdf':
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() or ''
        elif ext == 'docx':
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
    except Exception as e:
        return {"error": str(e)}

    ocr_time = round(time() - start, 2)

    classification = {"type": "Other", "confidence": 0, "reason": "Not classified"}
    try:
        response = classify_document(text)
        if response.strip().startswith("```"):
            response = response.strip("` \n").replace("```json", "").replace("```", "")
        classification = json.loads(response)
    except Exception as e:
        classification["reason"] = f"Gemini Error: {str(e)}"

    # 🔄 Update document info
    data = _load_data()
    for doc in data:
        if doc['path'] == file_path or os.path.basename(doc['path']) == os.path.basename(file_path):
            doc['ocr_text'] = text
            doc['ocr_time'] = ocr_time
            doc['type'] = classification.get("type", "Other")
            doc['confidence'] = classification.get("confidence", 0)
            doc['reason'] = classification.get("reason", "No reason provided.")
            doc['status'] = "Classified"
            doc['last_updated'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            break
    _save_data(data)

    return {"text": text, "ocr_time": ocr_time, "classification": classification}

# ✅ Load all documents
def load_documents():
    return _load_data()

# ✅ Fetch document by ID
def get_document_by_id(doc_id):
    for doc in _load_data():
        if doc['id'] == doc_id:
            return doc
    return None

# ✅ Fetch document by name (from data or log)
def get_document_by_name(filename):
    # Check in documents.json
    for doc in _load_data():
        if doc['name'] == filename:
            return doc

    # Fallback to document_log.json
    if os.path.exists(LOG_FILE):
        with open(LOG_FILE, "r") as f:
            logs = json.load(f)
        for log in logs:
            if log.get("filename") == filename:
                return {
                    "id": log.get("id", ""),
                    "name": log.get("filename"),
                    "path": os.path.join(UPLOAD_DIR, log.get("filename")),
                    "type": log.get("type", "Others"),
                    "status": log.get("status", "Classified"),
                    "last_updated": log.get("last_updated", "")
                }
    return None

# ✅ OCR Only (used for Re-Extract page)
def run_ocr(file_path):
    ext = file_path.split('.')[-1].lower()
    text = ""
    ocr_time = 0.0

    from time import time
    start = time()

    try:
        if ext in ['png', 'jpg', 'jpeg']:
            text = pytesseract.image_to_string(Image.open(file_path))
        elif ext == 'pdf':
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() or ''
        elif ext == 'docx':
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
    except Exception as e:
        return {"error": str(e), "text": "", "ocr_time": 0.0}

    ocr_time = round(time() - start, 2)

    # Update text in the saved document
    data = _load_data()
    for doc in data:
        if doc['path'] == file_path:
            doc['ocr_text'] = text
            doc['ocr_time'] = ocr_time
            doc['status'] = "OCR Extracted"
            doc['last_updated'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            break

    _save_data(data)
    return {"text": text, "ocr_time": ocr_time}
